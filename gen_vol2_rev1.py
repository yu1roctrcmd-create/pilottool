#!/usr/bin/env python3
"""NCA井戸端会議 Vol.2 Rev.1 Word document generator"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ======== PAGE SETUP ========
for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# ======== HELPER FUNCTIONS ========

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_regulation_box(doc, title_text, items, bg_hex='FFFFFF'):
    doc.add_paragraph('')
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg_hex)
    p = cell.paragraphs[0]
    if title_text:
        run = p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(9.5)
    for item in items:
        p2 = cell.add_paragraph(item)
        for run in p2.runs:
            run.font.size = Pt(9)
    doc.add_paragraph('')
    return table

def add_bubble(doc, speaker_label, text, bg_hex='FFD700', align=WD_ALIGN_PARAGRAPH.LEFT):
    """Add a speech bubble as a bordered, colored table cell."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg_hex)
    # Speaker label
    p_label = cell.paragraphs[0]
    p_label.alignment = align
    run_label = p_label.add_run(speaker_label)
    run_label.bold = True
    run_label.font.size = Pt(8.5)
    run_label.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    # Speech text
    p_text = cell.add_paragraph(text)
    p_text.alignment = align
    for run in p_text.runs:
        run.font.size = Pt(9.5)
    doc.add_paragraph('')
    return table

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f'■ {text}')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    doc.add_paragraph('')

def add_center_bold(doc, text, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)

def add_footer_table(doc, ref, issued_date, issued_by, prepared_by, pages):
    doc.add_paragraph('')
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    rows_data = [
        ['Reference #', ref, 'Issued date', issued_date],
        ['Issued by', issued_by, 'Prepared by', prepared_by],
        ['# of pages', pages, 'Crew Portal Site', 'Posted on "Announcement"'],
    ]
    for i, row_data in enumerate(rows_data):
        for j, text in enumerate(row_data):
            cell = table.cell(i, j)
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(8)
            if j % 2 == 0:
                run.bold = True

# ======================================================
#  PAGE 1 ― タイトル・シナリオ導入・日本のルール
# ======================================================

# --- タイトルヘッダー ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SHIBAYAMA SMALL TOWN TALK')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('NCA 井戸端会議')
run2.bold = True
run2.font.size = Pt(16)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('─' * 65)
run3.font.size = Pt(7)

p4 = doc.add_paragraph()
run_vol = p4.add_run('Vol. 2  Rev.1')
run_vol.bold = True
run_vol.font.size = Pt(12)
run_date = p4.add_run('                                                                     March 2, 2026')
run_date.font.size = Pt(11)

doc.add_paragraph('')

# --- メインタイトル ---
title_table = doc.add_table(rows=1, cols=1)
title_table.style = 'Table Grid'
tc = title_table.cell(0, 0)
set_cell_bg(tc, 'EBF3FB')
tp = tc.paragraphs[0]
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run('アプローチ中のスピードコントロールについて（総集編）')
tr.bold = True
tr.font.size = Pt(14)
doc.add_paragraph('')

# --- シナリオ導入 ---
p = doc.add_paragraph('　ある日、Bさんが香港へApproach中のこと…')
for run in p.runs:
    run.font.size = Pt(10)

add_regulation_box(doc, '',
    ['NCA203  Reduce Speed 160kts',
     'NCA203  Turn Left Direct RIVER  Cleared ILS RWY 25R APP',
     'NCA203  Contact Tower 118.2'],
    bg_hex='EBF3FB')

add_bubble(doc, '【Pilot B】',
    'あら？ どこまで160ktsを維持したらいいのかしら？\n'
    'Approach Clearanceが出たけど、速度指示はCancelになったの？\n'
    '日本だったら自動で終了するはずだけど… 香港はどうなの？ あれー？？？',
    bg_hex='FFD700')

add_bubble(doc, '【Pilot A】',
    'おぅ、Bさん！ ほな、まず日本の規程からおさらいして、\n'
    '香港（ICAO）とFAAも比べてみよかー！',
    bg_hex='90EE90',
    align=WD_ALIGN_PARAGRAPH.RIGHT)

# --- 日本のルール ---
add_section_heading(doc, '日本のルール（AIP JAPAN）')

add_regulation_box(doc, 'AIP JAPAN  ENR 1.8.6',
    ['次に掲げる場合は、それ以前に指示されていた速度調整は 自動的に終了する。',
     'a)  待機が指示された場合',
     'b)  「Climb via SID」又は「Descend via STAR」が指示された場合',
     'c)  進入許可が発出され、管制官から速度調整に係る指示が新しく又は繰り返して行われなかった場合',
     'd)  レーダー進入において接地点から5マイルの地点又は最終降下開始点のうち\n'
     '      いずれか接地点から遠い方の地点を通過したのち',
     'e)  速度を維持すべき地点が明示されたのち当該地点を通過した場合'],
    bg_hex='FFFFFF')

add_bubble(doc, '【Pilot B】',
    'AIP JAPAN ENR1.8.6のc)を見ると、Approach Clearanceが発出された時点で\n'
    '速度調整は自動的に終了ね。だからPilot\'s Discretionで減速できるわ。',
    bg_hex='FFD700')

add_bubble(doc, '【Pilot A】',
    'そうや！ 管制官がわざわざ「Speed Restriction Cancelled」って言わなくても\n'
    '自動的に終わるんやな。これは日本独自のルールやで。\n'
    'では、香港はどうやろ？',
    bg_hex='90EE90',
    align=WD_ALIGN_PARAGRAPH.RIGHT)

# --- フッター P.1 ---
add_footer_table(doc,
    ref='TYOOB 18-022-Rev1',
    issued_date='March 2, 2026',
    issued_by='TYOOBKZ',
    prepared_by='Hara  yosuke.hara@nca.aero',
    pages='4')

# ======================================================
#  PAGE 2 ― 香港（ICAO Doc 4444）
# ======================================================
doc.add_page_break()

add_section_heading(doc, '香港のルール（AIP HONG KONG / ICAO Doc 4444）')

add_regulation_box(doc, 'AIP HONG KONG  ENR 1 GENERAL RULES AND PROCEDURES  ENR 1.1 GENERAL RULES',
    ['1.1  The air traffic rules and procedures applicable to air traffic within the',
     '      Hong Kong FIR conform to: Annex 2 and Annex 11, the Air Navigation',
     '      (Hong Kong) Order 1995; ICAO Doc 4444 PANS/ATM, and the Regional',
     '      Supplementary Procedures MID/ASIA Region, except for the differences listed in GEN 1.7.'],
    bg_hex='FFFFFF')

add_bubble(doc, '【Pilot B】',
    'ふむふむ、香港はICAO Doc 4444に従うのね。\n'
    'ICAO Differencesを見たところ、速度調整に関する相違点はなさそうね。\n'
    'では、Doc 4444にはなんて書いてあるのかしら？',
    bg_hex='FFD700')

add_bubble(doc, '【Pilot A】',
    '勉強熱心やなー。 ほれ、これや。',
    bg_hex='90EE90',
    align=WD_ALIGN_PARAGRAPH.RIGHT)

add_regulation_box(doc, 'AIR TRAFFIC MANAGEMENT (DOC 4444)',
    ['4.6 HORIZONTAL SPEED CONTROL INSTRUCTIONS',
     '',
     '4.6.1.2  Speed control instructions shall remain in effect unless explicitly',
     '              cancelled or amended by the controller.',
     '',
     '4.6.3.7  Speed control should not be applied to aircraft after passing a point',
     '              7 km (4 NM) from the threshold on final approach.',
     '',
     'NOTE:  The flight crew has a requirement to fly a stabilized approach (airspeed and',
     '            configuration) typically by 5 km (3 NM) from the threshold',
     '            (Doc 8168, PANS-OPS, Volume I Part III Section 4, Chapter 3, 3.3 refers)'],
    bg_hex='FFFFFF')

add_bubble(doc, '【Pilot B】',
    '香港ではApproach Clearanceが発出されてもSpeed ControlはCancelにならない。\n'
    'Doc 4444から、4NMまでにspeedを変更する場合はリクエストベースね。\n'
    'Approach Speedを決める際は、Stabilized Approach、Weather、Trafficなどを\n'
    '考慮した総合的な判断が大事ということね。',
    bg_hex='FFD700')

add_bubble(doc, '【Pilot A】',
    'どう？ 納得？\n'
    'せやな。日本とは大きく違うポイントやな。\n'
    'では、次はFAAを見てみよか！',
    bg_hex='90EE90',
    align=WD_ALIGN_PARAGRAPH.RIGHT)

# --- フッター P.2 ---
add_footer_table(doc,
    ref='TYOOB 18-022-Rev1',
    issued_date='March 2, 2026',
    issued_by='TYOOBKZ',
    prepared_by='Hara  yosuke.hara@nca.aero',
    pages='4')

# ======================================================
#  PAGE 3 ― FAA（JO 7110.65）
# ======================================================
doc.add_page_break()

add_section_heading(doc, 'FAA（米国）のルール（FAA Order JO 7110.65）')

add_bubble(doc, '【Pilot C】',
    'ひょっこりはん！ FAA編、俺に任せてや！',
    bg_hex='FFA07A',
    align=WD_ALIGN_PARAGRAPH.RIGHT)

add_bubble(doc, '【Pilot A】',
    '待ってたで、Pilot C！ FAAはどうなんや？',
    bg_hex='90EE90')

add_bubble(doc, '【Pilot C】',
    'FAAのルールはJO 7110.65に規定されてるで。まず確認してみよか！',
    bg_hex='FFA07A',
    align=WD_ALIGN_PARAGRAPH.RIGHT)

add_regulation_box(doc, 'FAA Order JO 7110.65  Section 5-7-1  SPEED ADJUSTMENT',
    ['a.  Apply speed adjustment procedures to achieve or maintain required or desired spacing.',
     '',
     'b.  Do not assign speed adjustments to aircraft:',
     '      1.  At or inside the final approach fix (FAF) on final, or',
     '      2.  A point 5 miles from the runway,',
     '      whichever is closer to the runway.',
     '',
     'c.  Speed adjustment instructions shall remain in effect until explicitly cancelled or',
     '      modified by the controller.',
     '      ※ Approach Clearanceの発出のみでは速度指示はCancelされない。',
     '         管制官が明示的に「Cancel Speed Restrictions」等と指示した場合のみ終了する。'],
    bg_hex='FFFFFF')

add_bubble(doc, '【Pilot C】',
    'FAAもDoc 4444と同じく、Approach Clearanceだけでは速度指示はCancelにならへん。\n'
    '管制官が明示的にCancelと言うまで有効やで。',
    bg_hex='FFA07A',
    align=WD_ALIGN_PARAGRAPH.RIGHT)

add_bubble(doc, '【Pilot B】',
    'ICAOとの違いはどこにあるの？',
    bg_hex='FFD700')

add_bubble(doc, '【Pilot C】',
    'ポイントは速度管理の対象外となる地点やな。\n'
    'ICAOは Threshold から4NM以内、\n'
    'FAAは Runway から5NM以内、またはFAFの内側（滑走路に近い方）が対象外になるで。',
    bg_hex='FFA07A',
    align=WD_ALIGN_PARAGRAPH.RIGHT)

add_bubble(doc, '【Pilot B】',
    'なるほど！ FAAはFAFの位置がポイントになるのね。\n'
    'ILSアプローチではFAFはだいたい5〜7NM付近にあることが多いから、\n'
    'ケースによってはICAOより早く速度制限の対象外になることもあるわね。',
    bg_hex='FFD700')

add_bubble(doc, '【Pilot A】',
    'ええな！ 3地域が出そろったな。最後にまとめてみよか！',
    bg_hex='90EE90',
    align=WD_ALIGN_PARAGRAPH.RIGHT)

# --- フッター P.3 ---
add_footer_table(doc,
    ref='TYOOB 18-022-Rev1',
    issued_date='March 2, 2026',
    issued_by='TYOOBKZ',
    prepared_by='Hara  yosuke.hara@nca.aero',
    pages='4')

# ======================================================
#  PAGE 4 ― まとめ比較表
# ======================================================
doc.add_page_break()

add_section_heading(doc, 'まとめ：3地域の速度調整ルール比較')

add_bubble(doc, '【Pilot B】',
    '3地域をまとめてみましょう！',
    bg_hex='FFD700')

# --- 比較テーブル ---
doc.add_paragraph('')
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'

# ヘッダー行
headers = ['項目', '日本\n(AIP JAPAN)', '香港\n(ICAO Doc 4444)', 'FAA\n(JO 7110.65)']
hrow = table.rows[0]
for i, h in enumerate(headers):
    cell = hrow.cells[i]
    set_cell_bg(cell, '2E75B6')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(255, 255, 255)

# データ行
data = [
    ['Approach Clearanceで\n速度指示が自動終了',
     '○\n（自動終了）',
     '✗\n（自動終了しない）',
     '✗\n（自動終了しない）'],
    ['速度指示の\n終了タイミング',
     'Approach Clearance\n発出時（c項）',
     '管制官が明示的に\nCancelした時',
     '管制官が明示的に\nCancelした時'],
    ['速度管理の\n対象外となる地点',
     '接地点から5NM\nまたは最終降下開始点\n（遠い方）',
     'Thresholdから\n4NM以内',
     'Runwayから5NM以内\nまたはFAFの内側\n（滑走路に近い方）'],
    ['根拠規程',
     'AIP JAPAN\nENR 1.8.6',
     'ICAO Doc 4444\n4.6.1.2 / 4.6.3.7',
     'FAA Order\nJO 7110.65\nSection 5-7-1'],
]

row_bg_even = 'DEEAF1'
row_bg_odd = 'FFFFFF'
highlight_ja = 'FFE699'  # 日本の独自ルール部分を強調

for row_idx, row_data in enumerate(data):
    row = table.rows[row_idx + 1]
    for col_idx, text in enumerate(row_data):
        cell = row.cells[col_idx]
        # 背景色設定
        if col_idx == 0:
            set_cell_bg(cell, 'BDD7EE')
        elif row_idx == 0 and col_idx == 1:
            set_cell_bg(cell, highlight_ja)  # 日本が独自のところを黄色強調
        elif row_idx % 2 == 0:
            set_cell_bg(cell, row_bg_even)
        else:
            set_cell_bg(cell, row_bg_odd)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(9)
        if col_idx == 0:
            run.bold = True
        if row_idx == 0 and col_idx == 1:
            run.bold = True

doc.add_paragraph('')

# --- まとめの会話 ---
add_bubble(doc, '【Pilot A】',
    '日本は独自ルールで、Approach Clearanceで速度指示が自動終了するんやな。\n'
    'HKGもFAAも管制官の明示的なCancelが必要や。\n'
    '海外フライト時は必ずその国のAIPを確認することが大事やで！',
    bg_hex='90EE90',
    align=WD_ALIGN_PARAGRAPH.RIGHT)

add_bubble(doc, '【Pilot B】',
    '特に日本から海外に行ったときは、つい日本のルールで考えてしまいがち。\n'
    'Approach Clearance後も速度指示は生きていると意識して、\n'
    'しっかり確認しながら飛ぶことが大切ね！',
    bg_hex='FFD700')

add_bubble(doc, '【Pilot C】',
    'もうええわ、おおきに！ でも他の国も気になるな〜。また調べてみよか！',
    bg_hex='FFA07A',
    align=WD_ALIGN_PARAGRAPH.RIGHT)

# 次回予告
doc.add_paragraph('')
add_center_bold(doc, '好評につき、次回につづく！', size=13)
doc.add_paragraph('')

# --- フッター P.4 ---
add_footer_table(doc,
    ref='TYOOB 18-022-Rev1',
    issued_date='March 2, 2026',
    issued_by='TYOOBKZ',
    prepared_by='Hara  yosuke.hara@nca.aero',
    pages='4')

# ======== SAVE ========
output_path = '/Users/yuichiromori/Desktop/NCAフォルダ/NCA井戸端会議_VOL.2_Rev1.docx'
doc.save(output_path)
print(f'✓ Word document saved: {output_path}')
