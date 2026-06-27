#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
VOL_26 Conflict Zone周辺の飛行 - リバイス案（コンパクト版）生成スクリプト
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

def add_dialogue(doc, speaker, text, speaker_rgb):
    p = doc.add_paragraph()
    run_s = p.add_run(f"【{speaker}】  ")
    run_s.bold = True
    run_s.font.size = Pt(10.5)
    run_s.font.color.rgb = RGBColor(*speaker_rgb)
    run_t = p.add_run(text)
    run_t.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(5)
    return p

def add_section_title(doc, text, new=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(180, 0, 0) if new else RGBColor(60, 100, 180)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p

def shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_hex)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def set_cell_text(cell, text, size=9, bold=False, color=None, align=None):
    cell.text = text
    for para in cell.paragraphs:
        if align:
            para.alignment = align
        for run in para.runs:
            run.font.size = Pt(size)
            run.bold = bold
            if color:
                run.font.color.rgb = RGBColor(*color)

def add_separator(doc):
    p = doc.add_paragraph("─" * 65)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(7)
    p.runs[0].font.color.rgb = RGBColor(200, 200, 200)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

# =========================================
# ヘッダー
# =========================================
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = p_title.add_run("SHIBAYAMA SMALL TOWN TALK　NCA 井戸端会議\n")
r1.bold = True
r1.font.size = Pt(16)
r1.font.color.rgb = RGBColor(52, 100, 200)
r2 = p_title.add_run("Vol. 26  Rev.1　　April 2026")
r2.font.size = Pt(11)
r2.bold = True

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sub = p_sub.add_run("Conflict Zone 周辺の飛行")
r_sub.bold = True
r_sub.font.size = Pt(15)

# 改訂注記
p_rev = doc.add_paragraph()
p_rev.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_rev = p_rev.add_run(
    "【改訂概要】2026/4/1発効 OM S-3-11 にてConflict ZoneのCategory分類が新設。内容をアップデート。\n"
    "（OPS UPDATE No.26005 / OM S-3-11 Eff. 20260401）"
)
r_rev.font.size = Pt(9)
r_rev.font.color.rgb = RGBColor(180, 0, 0)

add_separator(doc)

# =========================================
# Part 1: Conflict Zone とは（コンパクト）
# =========================================
add_section_title(doc, "◆ Conflict Zone とは")

add_dialogue(doc, "Pilot C",
    "欧州パターンから帰ってきたわ。ロシア・ウクライナ戦争も終わらへんし、中東の紛争もあるし……世界は混沌としてるなぁ。",
    (60, 130, 60))

add_dialogue(doc, "Pilot B",
    "NCAでも日本-欧州間で迂回ルートを強いられ、黒海付近ではGPS Signal Interferencesもあるよね。",
    (60, 60, 180))

add_dialogue(doc, "Pilot A",
    "そういった軍事組織間の紛争地域を「Conflict Zone」という。OM 3-4-2にも定義があり、"
    "飛行計画時にConflict Zone周辺の安全情報を収集し、経路・巡航高度を適切に設定することが求められているよ。",
    (180, 100, 0))

# OM3-4-2（1行コンパクト）
p_om = doc.add_paragraph()
r_om = p_om.add_run(
    "OM 3-4-2：Conflict Zone とは、軍事組織間で武力衝突が発生している・発生する可能性がある、"
    "または軍事組織の警戒・緊張状態が高まっている地帯をいう。"
)
r_om.font.size = Pt(9)
r_om.font.color.rgb = RGBColor(130, 0, 0)
p_om.paragraph_format.left_indent = Cm(0.8)
p_om.paragraph_format.right_indent = Cm(0.8)
p_om.paragraph_format.space_after = Pt(6)

add_separator(doc)

# =========================================
# Part 2: カテゴリー分類（NEW・メイン）
# =========================================
add_section_title(doc, "◆ 【NEW】Conflict Zone のCategory分類（OM S-3-11 / 2026年4月1日発効）", new=True)

add_dialogue(doc, "Pilot A",
    "今年の4月から OM S-3-11 が発効して、Conflict Zone の扱いが大きく変わったよ。\n"
    "これまでは公的機関（FAA・EASA等）の評価をもとに原則回避していたけど、"
    "会社として独自にリスク評価を行い、Conflict ZoneをCategory 1〜3に分類して運航管理するようになったんだ。",
    (180, 100, 0))

add_dialogue(doc, "Pilot B",
    "カテゴリーごとに何が変わるの？",
    (60, 60, 180))

add_dialogue(doc, "Pilot C",
    "飛行計画段階と飛行中でそれぞれ制限内容が変わるよ。表にまとめると↓",
    (60, 130, 60))

# カテゴリー表
doc.add_paragraph()
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'

# ヘッダー
h_data = ["Category", "分類", "飛行計画段階", "飛行中"]
for i, h in enumerate(h_data):
    cell = table.rows[0].cells[i]
    set_cell_text(cell, h, size=9, bold=True, color=(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(cell, "2E5FA3")

# Cat 1
r1_data = [
    "Category 1", "Prohibited\n（飛行禁止）",
    "当該区域を回避\n（Flight prohibited）",
    "当該区域への入域不可\n（Entry prohibited）"
]
for i, text in enumerate(r1_data):
    cell = table.rows[1].cells[i]
    set_cell_text(cell, text, size=9)
    shade_cell(cell, "FFD7D7")

# Cat 2
r2_data = [
    "Category 2", "Restricted\n（飛行制限）",
    "当該区域内の指定高度を回避\n※運航上やむを得ず所定条件を満たす場合を除く",
    "指定高度を回避\n※飛行計画段階で確認した条件を満足する場合を除く"
]
for i, text in enumerate(r2_data):
    cell = table.rows[2].cells[i]
    set_cell_text(cell, text, size=9)
    shade_cell(cell, "FFF2CC")

# Cat 3
r3_data = [
    "Category 3", "Information\n（情報提供）",
    "最新情報に留意しつつ通常運航可",
    "最新情報に留意しつつ通常運航可"
]
for i, text in enumerate(r3_data):
    cell = table.rows[3].cells[i]
    set_cell_text(cell, text, size=9)
    shade_cell(cell, "E2EFDA")

# 制限なし
r4_data = [
    "制限なし", "—",
    "制限なし（公的機関情報をもとに管理）",
    "制限なし（COO指示のもと各部長が設定）"
]
for i, text in enumerate(r4_data):
    cell = table.rows[4].cells[i]
    set_cell_text(cell, text, size=9)
    shade_cell(cell, "F2F2F2")

doc.add_paragraph()

p_cat_note = doc.add_paragraph()
r_cn = p_cat_note.add_run(
    "※ カテゴリーはICAO Doc.10084ベースのリスク評価（脅威の性質・想定影響・緩和措置の有無等）で決定。\n"
    "   Flight Restriction Airspace List（TYOOE発行）で確認。問合せ：tyooekz@nca.aero"
)
r_cn.font.size = Pt(8.5)
r_cn.font.color.rgb = RGBColor(100, 100, 100)
p_cat_note.paragraph_format.left_indent = Cm(0.5)

add_dialogue(doc, "Pilot A",
    "適用範囲は計画経路から 50NM 以内のConflict Zone。\n"
    "乗務員と運航管理者は、少なくとも50NM以内に存在するConflict Zoneの情報を確認し、"
    "必要な情報共有・飛行支援を行うことが求められているよ。",
    (180, 100, 0))

add_dialogue(doc, "Pilot C",
    "フライト中に新たな事態が発生した場合は、当該FIRを確認し、"
    "継続か目的地変更かを機長と運航管理者が協議して決定することになっているよ。",
    (60, 130, 60))

add_separator(doc)

# =========================================
# Part 3: FD Pro 表示（NEW）
# =========================================
add_section_title(doc, "◆ 【NEW】FD Pro での Conflict Zone 表示", new=True)

add_dialogue(doc, "Pilot B",
    "Conflict Zoneって、FD Proでも見えるようになったって本当？",
    (60, 60, 180))

add_dialogue(doc, "Pilot A",
    "そう！ Enroute Chart上にカテゴリーごとに色分けで表示されるよ。\n"
    "  ・Category 1：茶色の点線＋塗りつぶし\n"
    "  ・Category 2：赤色の点線＋塗りつぶし\n"
    "  ・Category 3：茶色の点線のみ\n"
    "表示単位はFIR全体。BoundaryラインやBall Noteをタップすると詳細情報が確認できる\n"
    "（計画ルートが当該空域を通過している場合のみ）。",
    (180, 100, 0))

add_dialogue(doc, "Pilot C",
    "フライト前のブリーフィングで視覚的に確認できるのは助かるね。\n"
    "初回表示と情報更新のたびにアップデートが必要だから、最新状態を確認するようにしてね。",
    (60, 130, 60))

add_separator(doc)

# =========================================
# Part 4: 要撃・その他（コンパクト）
# =========================================
add_section_title(doc, "◆ 要撃（INTERCEPTION）・その他")

add_dialogue(doc, "Pilot B",
    "万が一要撃された場合は OM 10-4-13 の手順を迅速に実施。\n"
    "Airway Manual の視覚信号に従い、緊急周波数121.5MHzで通信・Transponder 7700をセット。\n"
    "NORADのTFR・Interception動画も参考になるよ。",
    (60, 60, 180))

add_dialogue(doc, "Pilot A",
    "軍用艦が空域を迂回させるために航空機へ飛行方向を指示した事例もある。\n"
    "（IFALPA: Communication Interference by Military Warships in the Pacific Region）\n"
    "当該空域付近ではDeviationや Direct Waypointリクエスト時にno-fly zoneを把握しておくことが重要だよ。",
    (180, 100, 0))

add_dialogue(doc, "Pilot A",
    "今日家に帰って休んだら色々とレビューせんとなぁ。\n"
    "妻と喧嘩したまま仕事に出てきて、家庭が「Do not fly!」レベルのConflict Zoneなのを忘れとった。\n"
    "要撃されたらSQ7700セットして妻の指示に大人しく従いますー！！",
    (180, 100, 0))

add_separator(doc)

# =========================================
# 参考文献
# =========================================
p_ref = doc.add_paragraph()
r_ref = p_ref.add_run("○ 参考文献")
r_ref.bold = True
r_ref.font.size = Pt(10)

refs = [
    ("OM 3-4-2 ⑦ Conflict Zone", False),
    ("OM S-3-11 Operations over and near Conflict Zones（Eff. 20260401）", True),
    ("OM 10-4-13 要撃（INTERCEPTION）", False),
    ("OR OMR 2. Conflict Zone", False),
    ("運航関連業務管理規則　Conflict Zoneに係る飛行空域運用要領", False),
    ("OPS UPDATE No.24008 Changes to the company routes around the Black Sea", False),
    ("OPS UPDATE No.26003 Operations of Flights Considering the Situation in the Middle East", True),
    ("OPS UPDATE No.26005 Explanation of OM S-3-11 'Operations over and near Conflict Zones'", True),
    ("IFALPA Communication Interference by Military Warships in the Pacific Region", False),
    ("IFALPA Flying into and Over Conflict Zones", False),
]

for text, is_new in refs:
    p = doc.add_paragraph()
    label = "  ★" if is_new else "  ・"
    r = p.add_run(f"{label}{text}")
    r.font.size = Pt(9)
    if is_new:
        r.font.color.rgb = RGBColor(180, 0, 0)
        r.bold = True
    p.paragraph_format.space_after = Pt(1)

# フッター
doc.add_paragraph()
p_f = doc.add_paragraph()
r_f = p_f.add_run(
    "Reference #: TYOO1 (Rev.1)　Issued date: April 2026　Prepared by: TYOO1　"
    "Total # of pages: TBD　Crew Portal Site: Posted on \"Announcement\""
)
r_f.font.size = Pt(8)
r_f.font.color.rgb = RGBColor(120, 120, 120)
p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER

output_path = (
    "/Users/yuichiromori/Library/Mobile Documents/com~apple~CloudDocs/"
    "1.NCA業務/22_芝山業務/Conflict/VOL_26_Conflict_Zone_Rev1_草案.docx"
)
doc.save(output_path)
print(f"保存完了: {output_path}")
